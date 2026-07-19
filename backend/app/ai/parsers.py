"""Text extraction from uploaded knowledge documents (PDF, DOCX, TXT)."""

from pathlib import Path

from app.models.document import DocumentType


class DocumentParseError(Exception):
    """Raised when a document cannot be parsed or contains no extractable text."""


def extract_text(path: Path, file_type: DocumentType) -> str:
    """Extract raw text from a stored document file.

    Raises DocumentParseError if the file is unreadable or yields no text.
    """
    try:
        if file_type is DocumentType.PDF:
            text = _extract_pdf(path)
        elif file_type is DocumentType.DOCX:
            text = _extract_docx(path)
        else:
            text = _extract_txt(path)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Could not parse {file_type.value} file.") from exc

    if not text.strip():
        raise DocumentParseError("Document contains no extractable text.")
    return text


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cell for cell in cells if cell))
    return "\n".join(parts)


def _extract_txt(path: Path) -> str:
    data = path.read_bytes()
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1")
